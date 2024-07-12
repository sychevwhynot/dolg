from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import get_user_model
from .models import Shtraf, Spisanie
from django.contrib.auth import authenticate, login
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth import logout
from .forms import ShtrafForm, DeductFineForm, FilterShtrafForm, FilterSpisanieForm
from django.contrib.auth.decorators import user_passes_test, login_required
from users.models import Otdel
from django.http import HttpResponseRedirect
from django.db.models import Sum
import docx
from io import BytesIO
from django.http import HttpResponse
from django.http import JsonResponse
from django.views.generic import View
from django.db.models import Q
from django.utils.timezone import now
import datetime
from django.utils import timezone



User = get_user_model()

def login_view(request):
    if request.method == 'POST':
        form = AuthenticationForm(request, request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(username=username, password=password)
            if user is not None:
                login(request, user)
                return redirect('shtraf:profile_list')
    else:
        form = AuthenticationForm()
    return render(request, 'login.html', {'form': form})

def logout_view(request):
    logout(request)
    return redirect('login')

def my_profile(request):
    user = request.user
    
    # Получаем все штрафы текущего пользователя
    shtrafy = Shtraf.objects.filter(user=user).order_by('-date')
    
    # Вычисляем общую сумму штрафов
    total_shtrafy = shtrafy.aggregate(total_sum=Sum('summa'))['total_sum'] or 0
    
    # Получаем все списания текущего пользователя
    spisaniya = Spisanie.objects.filter(user=user).order_by('-date')
    
    # Вычисляем общую сумму списаний
    total_deductions = spisaniya.aggregate(total_sum=Sum('summa'))['total_sum'] or 0
    
    # Вычисляем текущую задолженность
    dynamic_shtrafy = total_shtrafy - total_deductions
    
    # Получаем списания за текущий месяц
    current_month = now().month
    current_year = now().year
    spisanie_current_month = spisaniya.filter(created_at__month=current_month, created_at__year=current_year)
    total_spisanie_current_month = spisanie_current_month.aggregate(Sum('summa'))['summa__sum'] or 0
    
    context = {
        'user': user,
        'shtrafy': shtrafy,
        'spisaniya': spisaniya,
        'total_shtrafy': total_shtrafy,
        'dynamic_shtrafy': dynamic_shtrafy,
        'total_spisanie_current_month': total_spisanie_current_month,
    }

    return render(request, 'shtraf/my_profile.html', context)


@login_required
def profile_detail(request, pk):
    user = get_object_or_404(User, pk=pk)
    shtrafy = Shtraf.objects.filter(user=user).order_by('-date')
    total_shtrafy = sum(shtraf.summa for shtraf in shtrafy)
    

    spisaniya = Spisanie.objects.filter(user=user).order_by('-date')
    total_deductions = sum(spisanie.summa for spisanie in spisaniya)
    dynamic_shtrafy = total_shtrafy - total_deductions

    current_month = timezone.now().month
    current_year = timezone.now().year
    spisanie_current_month = Spisanie.objects.filter(user=user, created_at__month=current_month, created_at__year=current_year)
    total_spisanie_current_month = spisanie_current_month.aggregate(Sum('summa'))['summa__sum'] or 0

    if request.method == 'POST' and request.user.is_buh:
        form = DeductFineForm(request.POST)
        if form.is_valid():
            spisanie = form.save(commit=False)
            spisanie.user = user
            spisanie.date = now().date()
            spisanie.save()
            

            total_deductions = sum(spisanie.summa for spisanie in Spisanie.objects.filter(user=user))
            dynamic_shtrafy = total_shtrafy - total_deductions
            
            return HttpResponseRedirect(request.path_info)
    else:
        form = DeductFineForm()



    context = {
        'user': user,
        'shtrafy': shtrafy,
        'total_shtrafy': total_shtrafy,
        'dynamic_shtrafy': dynamic_shtrafy,
        'form': form,
        'spisaniya': spisaniya,  # Добавляем списания в контекст
        'total_spisanie_current_month': total_spisanie_current_month,  # Добавляем сумму списаний за текущий месяц
    }

    return render(request, 'shtraf/profile_detail.html', context)

@login_required
def profile_list(request):
    users = User.objects.all().order_by('last_name')
    return render(request, 'shtraf/profile_list.html', {'users': users})

def is_admin(user):
    return user.is_admin

@login_required
@user_passes_test(is_admin)
def create_shtraf(request, pk):
    user = User.objects.get(pk=pk)
    
    if request.method == 'POST':
        form = ShtrafForm(request.POST)
        if form.is_valid():
            shtraf = form.save(commit=False)
            shtraf.user = user
            shtraf.save()
            return redirect('shtraf:profile_detail', pk=pk)
    else:
        form = ShtrafForm(initial={'user': user})  # Заполняем форму автоматически пользователем

    context = {
        'form': form,
        'user': user,
    }
    return render(request, 'shtraf/create_shtraf.html', context)

@login_required
def otdel_list(request):
    otdel = Otdel.objects.all().order_by('name')
    return render(request, 'shtraf/otdel_list.html', {'otdel': otdel})

@login_required
def otdel_detail(request, pk):
    otdel = get_object_or_404(Otdel, id=pk)
    users = User.objects.filter(otdel=otdel)

    total_shtrafy_sum_all = 0
    total_dynamic_shtrafy_sum_all = 0
    user_data = []
    all_shtrafy = []

    for user in users:
        shtrafy = Shtraf.objects.filter(user=user).order_by('-date')
        total_shtrafy = shtrafy.aggregate(total_sum=Sum('summa'))['total_sum'] or 0
        all_shtrafy.extend(list(shtrafy))
        
        spisaniya = Spisanie.objects.filter(user=user)
        total_deductions = spisaniya.aggregate(total_sum=Sum('summa'))['total_sum'] or 0
        dynamic_shtrafy = total_shtrafy - total_deductions
        
        user_data.append({
            'user': user,
            'total_shtrafy': total_shtrafy,
            'dynamic_shtrafy': dynamic_shtrafy,
            'total_spisanie': total_deductions,
        })

        total_shtrafy_sum_all += total_shtrafy
        total_dynamic_shtrafy_sum_all += dynamic_shtrafy

    context = {
        'otdel': otdel,
        'user_data': user_data,
        'total_shtrafy_sum_all': total_shtrafy_sum_all,
        'total_dynamic_shtrafy_sum_all': total_dynamic_shtrafy_sum_all,
        'all_shtrafy': sorted(all_shtrafy, key=lambda x: x.date, reverse=True),
        'can_view_shtrafs': request.user.is_glav or request.user.is_buh or request.user.is_admin or (request.user.is_leed and request.user.otdel == otdel),
    }

    return render(request, 'shtraf/otdel_detail.html', context)

def is_glav(user):
    return user.is_authenticated and user.is_glav

@login_required
@user_passes_test(is_glav)
def report_view(request):
    shtrafy = Shtraf.objects.all().order_by('date')
    spisaniya = Spisanie.objects.all().order_by('date')
    
    total_shtrafy_sum = shtrafy.aggregate(Sum('summa'))['summa__sum'] or 0
    total_dynamic_shtrafy_sum = total_shtrafy_sum - sum(user.deducted_amount for user in User.objects.all())
    total_spisanie_sum = spisaniya.aggregate(Sum('summa'))['summa__sum'] or 0
    
    current_month = timezone.now().month
    current_year = timezone.now().year
    spisanie_current_month = spisaniya.filter(date__month=current_month, date__year=current_year)
    total_spisanie_current_month = spisanie_current_month.aggregate(Sum('summa'))['summa__sum'] or 0

    selected_month = None
    sum_for_selected_month = None
    selected_period_sum = None

    if request.method == 'POST':
        form = FilterShtrafForm(request.POST)
        if form.is_valid():
            start_date = form.cleaned_data.get('start_date')
            end_date = form.cleaned_data.get('end_date')
            otdel = form.cleaned_data.get('otdel')
            month = form.cleaned_data.get('month')

            if start_date:
                shtrafy = shtrafy.filter(date__gte=start_date)
                spisaniya = spisaniya.filter(date__gte=start_date)
            if end_date:
                shtrafy = shtrafy.filter(date__lte=end_date)
                spisaniya = spisaniya.filter(date__lte=end_date)
            if otdel and otdel != 'Все отделы':
                shtrafy = shtrafy.filter(user__otdel=otdel)
                spisaniya = spisaniya.filter(user__otdel=otdel)
            if month:
                selected_month = month
                shtrafy = shtrafy.filter(date__month=month)
                spisaniya = spisaniya.filter(date__month=month)
                sum_for_selected_month = spisaniya.aggregate(month_sum=Sum('summa'))['month_sum'] or 0

            if start_date or end_date:
                selected_period_sum = spisaniya.aggregate(period_sum=Sum('summa'))['period_sum'] or 0

            total_shtrafy_sum = shtrafy.aggregate(Sum('summa'))['summa__sum'] or 0
            filtered_users = User.objects.filter(shtraf__in=shtrafy).distinct()
            total_dynamic_shtrafy_sum = total_shtrafy_sum - sum(user.deducted_amount for user in filtered_users)
            total_spisanie_sum = spisaniya.aggregate(Sum('summa'))['summa__sum'] or 0

            if 'generate_word' in request.POST:
                return generate_word_report(shtrafy, total_shtrafy_sum, total_dynamic_shtrafy_sum, spisaniya, total_spisanie_sum, total_spisanie_current_month, selected_month, sum_for_selected_month, selected_period_sum)

    else:
        form = FilterShtrafForm()

    context = {
        'shtrafy': shtrafy,
        'spisaniya': spisaniya,
        'form': form,
        'total_shtrafy_sum': total_shtrafy_sum,
        'total_dynamic_shtrafy_sum': total_dynamic_shtrafy_sum,
        'total_spisanie_sum': total_spisanie_sum,
        'total_spisanie_current_month': total_spisanie_current_month,
        'current_month': current_month,
        'current_year': current_year,
        'selected_month': selected_month,
        'sum_for_selected_month': sum_for_selected_month,
        'selected_period_sum': selected_period_sum,
    }

    return render(request, 'shtraf/report_view.html', context)

def is_leed(user):
    return user.is_authenticated and user.is_leed


@login_required
@user_passes_test(is_leed)
def leed_report_view(request):
    user_otdel = request.user.otdel  # Определяем отделение текущего пользователя
    shtrafy = Shtraf.objects.all().order_by('date')
    spisaniya = Spisanie.objects.all().order_by('date')
    
    total_shtrafy_sum = shtrafy.aggregate(Sum('summa'))['summa__sum'] or 0
    total_dynamic_shtrafy_sum = total_shtrafy_sum - sum(user.deducted_amount for user in User.objects.all())
    total_spisanie_sum = spisaniya.aggregate(Sum('summa'))['summa__sum'] or 0
    
    current_month = timezone.now().month
    current_year = timezone.now().year
    spisanie_current_month = spisaniya.filter(date__month=current_month, date__year=current_year)
    total_spisanie_current_month = spisanie_current_month.aggregate(Sum('summa'))['summa__sum'] or 0

    selected_month = None
    sum_for_selected_month = None
    selected_period_sum = None

    if request.method == 'POST':
        form = FilterShtrafForm(request.POST)
        if form.is_valid():
            start_date = form.cleaned_data.get('start_date')
            end_date = form.cleaned_data.get('end_date')
            month = form.cleaned_data.get('month')

            if start_date:
                shtrafy = shtrafy.filter(date__gte=start_date)
                spisaniya = spisaniya.filter(date__gte=start_date)
            if end_date:
                shtrafy = shtrafy.filter(date__lte=end_date)
                spisaniya = spisaniya.filter(date__lte=end_date)
            if month:
                selected_month = month
                shtrafy = shtrafy.filter(date__month=month)
                spisaniya = spisaniya.filter(date__month=month)
                sum_for_selected_month = spisaniya.aggregate(month_sum=Sum('summa'))['month_sum'] or 0

            if start_date or end_date:
                selected_period_sum = spisaniya.aggregate(period_sum=Sum('summa'))['period_sum'] or 0

            total_shtrafy_sum = shtrafy.aggregate(Sum('summa'))['summa__sum'] or 0
            filtered_users = User.objects.filter(shtraf__in=shtrafy).distinct()
            total_dynamic_shtrafy_sum = total_shtrafy_sum - sum(user.deducted_amount for user in filtered_users)
            total_spisanie_sum = spisaniya.aggregate(Sum('summa'))['summa__sum'] or 0

            if 'generate_word' in request.POST:
                return generate_word_report(shtrafy, total_shtrafy_sum, total_dynamic_shtrafy_sum, spisaniya, total_spisanie_sum, total_spisanie_current_month, selected_month, sum_for_selected_month, selected_period_sum)

    else:
        initial_data = {'otdel': user_otdel}  # Задаем начальные данные для формы
        form = FilterShtrafForm(initial=initial_data)  # Инициализируем форму с начальными данными

    context = {
        'shtrafy': shtrafy,
        'spisaniya': spisaniya,
        'form': form,
        'total_shtrafy_sum': total_shtrafy_sum,
        'total_dynamic_shtrafy_sum': total_dynamic_shtrafy_sum,
        'total_spisanie_sum': total_spisanie_sum,
        'total_spisanie_current_month': total_spisanie_current_month,
        'current_month': current_month,
        'current_year': current_year,
        'selected_month': selected_month,
        'sum_for_selected_month': sum_for_selected_month,
        'selected_period_sum': selected_period_sum,
    }

    return render(request, 'shtraf/leed_report_view.html', context)


@login_required
@user_passes_test(lambda u: u.is_buh)
def buh_report_view(request):
    spisanie_list = Spisanie.objects.all().order_by('date')
    total_spisanie_sum = spisanie_list.aggregate(total_sum=Sum('summa'))['total_sum'] or 0

    # Вычисление суммы списаний за текущий месяц
    current_month = timezone.now().month
    current_year = timezone.now().year
    spisanie_current_month = spisanie_list.filter(date__month=current_month, date__year=current_year)
    total_spisanie_current_month = spisanie_current_month.aggregate(Sum('summa'))['summa__sum'] or 0

    selected_month = None
    sum_for_selected_month = None
    selected_period_sum = None

    if request.method == 'POST':
        form = FilterSpisanieForm(request.POST)
        if form.is_valid():
            start_date = form.cleaned_data.get('start_date')
            end_date = form.cleaned_data.get('end_date')
            otdel = form.cleaned_data.get('otdel')
            month = form.cleaned_data.get('month')

            if start_date:
                spisanie_list = spisanie_list.filter(date__gte=start_date)
            if end_date:
                spisanie_list = spisanie_list.filter(date__lte=end_date)
            if otdel and otdel != 'Все отделы':
                spisanie_list = spisanie_list.filter(user__otdel=otdel)
            if month:
                selected_month = month
                spisanie_list = spisanie_list.filter(date__month=month)
                sum_for_selected_month = spisanie_list.aggregate(month_sum=Sum('summa'))['month_sum'] or 0

            if start_date or end_date:
                selected_period_sum = spisanie_list.aggregate(period_sum=Sum('summa'))['period_sum'] or 0

            if 'generate_word' in request.POST:
                return generate_spisanie_word_report(spisanie_list, total_spisanie_sum, total_spisanie_current_month, sum_for_selected_month, selected_period_sum, current_month, current_year, selected_month)

    else:
        form = FilterSpisanieForm()

    context = {
        'spisanie_list': spisanie_list,
        'form': form,
        'total_spisanie_sum': total_spisanie_sum,
        'total_spisanie_current_month': total_spisanie_current_month,
        'current_month': current_month,
        'current_year': current_year,
        'selected_month': selected_month,
        'sum_for_selected_month': sum_for_selected_month,
        'selected_period_sum': selected_period_sum,
    }

    return render(request, 'shtraf/buh_report_view.html', context)


@login_required
@user_passes_test(is_admin)
def shtraf_report_view(request):
    shtrafy = Shtraf.objects.all().order_by('date')
    total_shtrafy_sum = shtrafy.aggregate(Sum('summa'))['summa__sum'] or 0
    total_dynamic_shtrafy_sum = total_shtrafy_sum - sum(user.deducted_amount for user in User.objects.all())

    if request.method == 'POST':
        form = FilterShtrafForm(request.POST)
        if form.is_valid():
            start_date = form.cleaned_data.get('start_date')
            end_date = form.cleaned_data.get('end_date')
            otdel = form.cleaned_data.get('otdel')

            if start_date:
                shtrafy = shtrafy.filter(date__gte=start_date)
            if end_date:
                shtrafy = shtrafy.filter(date__lte=end_date)
            if otdel:
                shtrafy = shtrafy.filter(user__otdel=otdel)

            total_shtrafy_sum = shtrafy.aggregate(Sum('summa'))['summa__sum'] or 0
            filtered_users = User.objects.filter(shtraf__in=shtrafy).distinct()
            total_dynamic_shtrafy_sum = total_shtrafy_sum - sum(user.deducted_amount for user in filtered_users)

            if 'generate_word' in request.POST:
                return generate_shtraf_word_report(shtrafy, total_shtrafy_sum, total_dynamic_shtrafy_sum)
    else:
        form = FilterShtrafForm()

    context = {
        'shtrafy': shtrafy,
        'form': form,
        'total_shtrafy_sum': total_shtrafy_sum,
        'total_dynamic_shtrafy_sum': total_dynamic_shtrafy_sum,
    }

    return render(request, 'shtraf/shtraf_report_view.html', context)



def generate_shtraf_word_report(shtrafy, total_shtrafy_sum, total_dynamic_shtrafy_sum):
    document = docx.Document()
    document.add_heading('Отчет по штрафам', 0)

    document.add_heading('Список штрафов', level=1)
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Дата'
    hdr_cells[1].text = 'Сотрудник'
    hdr_cells[2].text = 'Сумма'
    hdr_cells[3].text = 'Причина'

    for shtraf in shtrafy:
        row_cells = table.add_row().cells
        row_cells[0].text = str(shtraf.date)
        row_cells[1].text = f"{shtraf.user.last_name} {shtraf.user.first_name} {shtraf.user.superlast_name}"
        row_cells[2].text = str(shtraf.summa)
        row_cells[3].text = shtraf.reason

    document.add_heading('Итоги', level=1)
    document.add_paragraph(f'Общая сумма штрафов: {total_shtrafy_sum} руб.')
    document.add_paragraph(f'Текущая задолженность по штрафам: {total_dynamic_shtrafy_sum} руб.')

    f = BytesIO()
    document.save(f)
    f.seek(0)

    response = HttpResponse(f.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=shtraf_report.docx'
    return response


def generate_word_report(shtrafy, total_shtrafy_sum, total_dynamic_shtrafy_sum, spisaniya, total_spisanie_sum, total_spisanie_current_month, selected_month=None, sum_for_selected_month=None, selected_period_sum=None):
    document = docx.Document()
    document.add_heading('Отчет по штрафам и списаниям', 0)

    # Список штрафов
    document.add_heading('Список штрафов', level=1)
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Дата'
    hdr_cells[1].text = 'Сотрудник'
    hdr_cells[2].text = 'Сумма'
    hdr_cells[3].text = 'Причина'

    for shtraf in shtrafy:
        row_cells = table.add_row().cells
        row_cells[0].text = str(shtraf.date)
        row_cells[1].text = f"{shtraf.user.last_name} {shtraf.user.first_name} {shtraf.user.superlast_name}"
        row_cells[2].text = str(shtraf.summa)
        row_cells[3].text = shtraf.reason

    # Список списаний
    document.add_heading('Список списаний', level=1)
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Дата'
    hdr_cells[1].text = 'Сотрудник'
    hdr_cells[2].text = 'Сумма'

    for spisanie in spisaniya:
        row_cells = table.add_row().cells
        row_cells[0].text = str(spisanie.date)
        row_cells[1].text = f"{spisanie.user.last_name} {spisanie.user.first_name} {spisanie.user.superlast_name}"
        row_cells[2].text = str(spisanie.summa)

    # Итоги по штрафам
    document.add_heading('Итоги по штрафам', level=1)
    document.add_paragraph(f'Общая сумма штрафов: {total_shtrafy_sum}')
    document.add_paragraph(f'Текущая задолженность по штрафам: {total_dynamic_shtrafy_sum}')

    # Итоги по списаниям
    document.add_heading('Итоги по списаниям', level=1)
    document.add_paragraph(f'Общая сумма списаний: {total_spisanie_sum}')
    document.add_paragraph(f'Сумма списаний за текущий месяц: {total_spisanie_current_month}')

    if selected_month:
        document.add_paragraph(f'Сумма списаний за выбранный месяц ({selected_month}): {sum_for_selected_month}')

    if selected_period_sum:
        document.add_paragraph(f'Сумма списаний за выбранный период: {selected_period_sum}')

    f = BytesIO()
    document.save(f)
    f.seek(0)

    response = HttpResponse(f.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=report.docx'
    return response

class DoctorAutocompleteView(View):
    def get(self, request, *args, **kwargs):
        if request.is_ajax():
            q = request.GET.get('q', '')
            users = User.objects.filter(
                Q(first_name__icontains=q) | Q(last_name__icontains=q)
            )
            results = [{'id': user.id, 'text': f'{user.first_name} {user.last_name}'} for user in users]
            return JsonResponse(results, safe=False)
        return JsonResponse({'error': 'Ajax request required'}, status=400)
    
def generate_spisanie_word_report(spisanie_list, total_spisanie_sum, total_spisanie_current_month, sum_for_selected_month, selected_period_sum, current_month, current_year, selected_month=None):
    document = docx.Document()
    document.add_heading('Отчет по списаниям', 0)

    document.add_heading('Список списаний', level=1)
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Дата'
    hdr_cells[1].text = 'Сотрудник'
    hdr_cells[2].text = 'Отдел'
    hdr_cells[3].text = 'Сумма'

    for spisanie in spisanie_list:
        row_cells = table.add_row().cells
        row_cells[0].text = str(spisanie.date)
        row_cells[1].text = f"{spisanie.user.last_name} {spisanie.user.first_name} {spisanie.user.superlast_name}"
        row_cells[2].text = spisanie.user.otdel.name if spisanie.user.otdel else ''
        row_cells[3].text = str(spisanie.summa)

    document.add_heading('Итоги', level=1)
    document.add_paragraph(f'Общая сумма списаний за все время: {total_spisanie_sum}')

    if selected_period_sum is not None:
        document.add_paragraph(f'Сумма списаний за выбранный период: {selected_period_sum}')
    elif sum_for_selected_month is not None:
        document.add_paragraph(f'Сумма списаний за {selected_month}/{current_year}: {sum_for_selected_month}')
    else:
        document.add_paragraph(f'Сумма списаний за текущий месяц ({current_month}/{current_year}): {total_spisanie_current_month}')

    f = BytesIO()
    document.save(f)
    f.seek(0)

    response = HttpResponse(f.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=spisanie_report.docx'
    return response
